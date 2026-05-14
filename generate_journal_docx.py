#!/usr/bin/env python3
"""
Générateur du journal «À l'Heure» en format DOCX (Word) — 8 pages A4
Usage : python3 generate_journal_docx.py [--date 2026-01-29] [--numero 1]
Output: public/journal/alheure-YYYY-MM-DD.docx
"""

import json, argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = Path(__file__).parent
BLUE   = RGBColor(14,  43,  98)
RED    = RGBColor(200, 48,  42)
WHITE  = RGBColor(255, 255, 255)
LGRAY  = RGBColor(245, 243, 238)
INK    = RGBColor(26,  26,  26)
INK2   = RGBColor(80,  80,  80)
INK3   = RGBColor(140, 140, 140)
GOLD   = RGBColor(180, 140, 30)

RUBRIQUE_LABELS = {
    "senegal":"SÉNÉGAL","afrique":"AFRIQUE","monde":"MONDE",
    "politique":"POLITIQUE","economie":"ÉCONOMIE","societe":"SOCIÉTÉ",
    "sport":"SPORT","culture":"CULTURE","diaspora":"DIASPORA",
}
BADGE_LABELS = {"rep":"GRAND REPORTAGE","longformat":"LONG FORMAT","video":"VIDÉO","live":"DIRECT"}

def fr_date(d: datetime) -> str:
    mois  = ["janvier","février","mars","avril","mai","juin","juillet","août","septembre","octobre","novembre","décembre"]
    jours = ["lundi","mardi","mercredi","jeudi","vendredi","samedi","dimanche"]
    return f"{jours[d.weekday()]} {d.day} {mois[d.month-1]} {d.year}"

def rub_label(a: dict) -> str:
    return BADGE_LABELS.get(a.get("badge",""), RUBRIQUE_LABELS.get(a.get("rubrique",""), a.get("rubrique","").upper()))

def body_text(body: list) -> str:
    return " ".join(b.get("text","") for b in body if b.get("type") in ("p","blockquote","pullquote"))

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top",top),("bottom",bottom),("left",left),("right",right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def set_col_widths(table, widths_cm: list):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW  = OxmlElement("w:tcW")
                twips = int(widths_cm[i] * 567)
                tcW.set(qn("w:w"), str(twips))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)

def remove_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    tblBd = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBd.append(el)
    tblPr.append(tblBd)

def add_colored_border_top(table, hex_color: str, size=12):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    tblBd = tblPr.find(qn("w:tblBorders"))
    if tblBd is None:
        tblBd = OxmlElement("w:tblBorders"); tblPr.append(tblBd)
    top = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    str(size))
    top.set(qn("w:color"), hex_color)
    tblBd.append(top)

def add_insideV_border(table, hex_color="C0B8A8", size=4):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    tblBd = tblPr.find(qn("w:tblBorders"))
    if tblBd is None:
        tblBd = OxmlElement("w:tblBorders"); tblPr.append(tblBd)
    iv = OxmlElement("w:insideV")
    iv.set(qn("w:val"),   "single")
    iv.set(qn("w:sz"),    str(size))
    iv.set(qn("w:color"), hex_color)
    tblBd.append(iv)

def add_rule(cell_or_doc, color_hex="C0B8A8", size=4):
    p   = cell_or_doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBd = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:color"), color_hex)
    pBd.append(bot); pPr.append(pBd)

def add_red_tag(cell, label: str):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(label)
    r.bold = True; r.font.size = Pt(6.5)
    r.font.color.rgb = RED; r.font.name = "Arial"

def add_article(cell, art: dict, title_size=13, show_img=True, img_w_cm=5.5, img_h_cm=3.2):
    add_red_tag(cell, rub_label(art))
    if show_img:
        url = art.get("imgUrl","")
        path = BASE / "public" / url.lstrip("/")
        if not path.exists():
            jpg = path.with_suffix(".jpg")
            if jpg.exists(): path = jpg
        if path.exists() and path.suffix.lower() != ".svg":
            try:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(4)
                run = p.add_run()
                run.add_picture(str(path), width=Cm(img_w_cm), height=Cm(img_h_cm))
            except Exception:
                pass
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(title_size * 1.2)
    r = p.add_run(art.get("title","Sans titre"))
    r.bold = True; r.font.size = Pt(title_size)
    r.font.color.rgb = INK; r.font.name = "Times New Roman"
    if art.get("dek"):
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(2)
        p2.paragraph_format.line_spacing = Pt(13)
        r2 = p2.add_run(art["dek"])
        r2.italic = True; r2.font.size = Pt(9.5)
        r2.font.color.rgb = INK2; r2.font.name = "Times New Roman"
    if art.get("author"):
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after  = Pt(2)
        r3 = p3.add_run(f"Par {art['author']}  ·  {art.get('date','')}")
        r3.bold = True; r3.font.size = Pt(6.5)
        r3.font.color.rgb = INK3; r3.font.name = "Arial"
    add_rule(cell)
    txt = body_text(art.get("body",[]))
    if txt:
        p4 = cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p4.paragraph_format.space_before = Pt(0)
        p4.paragraph_format.space_after  = Pt(0)
        p4.paragraph_format.line_spacing = Pt(12)
        r4 = p4.add_run(txt)
        r4.font.size = Pt(9); r4.font.color.rgb = INK
        r4.font.name = "Times New Roman"

def add_une_header(doc, edition_date: datetime, numero: int):
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    set_cell_bg(tbl.cell(0,0), "0E2B62")
    set_cell_margins(tbl.cell(0,0), top=160, bottom=80, left=200, right=200)
    c = tbl.cell(0,0)
    p1 = c.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after  = Pt(2)
    r1 = p1.add_run("À l'Heure")
    r1.bold = True; r1.italic = True
    r1.font.size = Pt(36); r1.font.color.rgb = WHITE
    r1.font.name = "Times New Roman"
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run("LE QUOTIDIEN D'INFORMATION")
    r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(168, 180, 207)
    r2.font.name = "Arial"
    tbl2 = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl2)
    add_colored_border_top(tbl2, "C8302A", size=16)
    set_cell_bg(tbl2.cell(0,0), "F5F3EE")
    set_cell_margins(tbl2.cell(0,0), top=60, bottom=60, left=200, right=200)
    info = (f"{fr_date(edition_date).upper()}   |   N° {numero:03d}   |   "
            f"PRIX : 100 FCFA   |   WWW.ALHEURE.SN   |   ISSN : 3020-0001")
    p3 = tbl2.cell(0,0).add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(0)
    r3 = p3.add_run(info)
    r3.bold = True; r3.font.size = Pt(7)
    r3.font.color.rgb = BLUE; r3.font.name = "Arial"
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_inner_header(doc, rubrique_label: str, edition_date: datetime, numero: int):
    tbl = doc.add_table(rows=1, cols=3)
    remove_table_borders(tbl)
    add_colored_border_top(tbl, "C8302A", size=12)
    for i in range(3):
        set_cell_bg(tbl.cell(0,i), "0E2B62")
    set_cell_margins(tbl.cell(0,0), top=80, bottom=80, left=150, right=80)
    set_cell_margins(tbl.cell(0,1), top=80, bottom=80, left=80,  right=80)
    set_cell_margins(tbl.cell(0,2), top=80, bottom=80, left=80,  right=150)
    set_col_widths(tbl, [4.5, 9.5, 4.5])
    p0 = tbl.cell(0,0).add_paragraph()
    p0.paragraph_format.space_before = Pt(0); p0.paragraph_format.space_after = Pt(0)
    r0 = p0.add_run("À l'Heure")
    r0.bold = True; r0.italic = True
    r0.font.size = Pt(13); r0.font.color.rgb = WHITE; r0.font.name = "Times New Roman"
    p1 = tbl.cell(0,1).add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0); p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run(rubrique_label)
    r1.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = RED; r1.font.name = "Arial"
    p2 = tbl.cell(0,2).add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(f"N° {numero:03d} · {fr_date(edition_date)}")
    r2.font.size = Pt(7); r2.font.color.rgb = RGBColor(168,180,207); r2.font.name = "Arial"
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def add_footer_line(doc, page_num: int, edition_date: datetime):
    tbl = doc.add_table(rows=1, cols=3)
    remove_table_borders(tbl)
    add_colored_border_top(tbl, "C8302A", size=8)
    for i in range(3):
        set_cell_bg(tbl.cell(0,i), "0E2B62")
    set_col_widths(tbl, [6.0, 6.5, 6.0])
    set_cell_margins(tbl.cell(0,0), top=60,bottom=60,left=150,right=80)
    set_cell_margins(tbl.cell(0,1), top=60,bottom=60,left=80, right=80)
    set_cell_margins(tbl.cell(0,2), top=60,bottom=60,left=80, right=150)
    def fp(cell, txt, al=WD_ALIGN_PARAGRAPH.LEFT):
        p = cell.add_paragraph(); p.alignment = al
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt); r.font.size = Pt(6.5)
        r.font.color.rgb = RGBColor(168,180,207); r.font.name = "Arial"
    fp(tbl.cell(0,0), "RIGUEUR · INDÉPENDANCE · PROXIMITÉ · INNOVATION")
    fp(tbl.cell(0,1), "alheure.info", WD_ALIGN_PARAGRAPH.CENTER)
    fp(tbl.cell(0,2), f"Page {page_num}", WD_ALIGN_PARAGRAPH.RIGHT)

def add_section_break(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pageBreakBefore")
    pb.set(qn("w:val"), "true")
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

def build_une(doc, articles: list, edition_date: datetime, numero: int):
    add_une_header(doc, edition_date, numero)
    vedette     = next((a for a in articles if a.get("featured")), articles[0] if articles else None)
    secondaires = [a for a in articles if a is not vedette][:3]
    if vedette:
        tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(tbl)
        set_col_widths(tbl, [10.5, 7.5])
        set_cell_bg(tbl.cell(0,0), "0E2B62")
        set_cell_margins(tbl.cell(0,0), top=120, bottom=120, left=200, right=160)
        set_cell_margins(tbl.cell(0,1), top=0,   bottom=0,   left=0,   right=0)
        tbl.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cl = tbl.cell(0,0)
        add_red_tag(cl, rub_label(vedette))
        p = cl.add_paragraph()
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(38)
        r = p.add_run(vedette.get("title",""))
        r.bold = True; r.font.size = Pt(32)
        r.font.color.rgb = WHITE; r.font.name = "Times New Roman"
        if vedette.get("dek"):
            p2 = cl.add_paragraph()
            p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(4)
            p2.paragraph_format.line_spacing = Pt(16)
            r2 = p2.add_run(vedette["dek"])
            r2.italic = True; r2.font.size = Pt(12)
            r2.font.color.rgb = WHITE; r2.font.name = "Times New Roman"
        if vedette.get("author"):
            p3 = cl.add_paragraph()
            p3.paragraph_format.space_before = Pt(2); p3.paragraph_format.space_after = Pt(4)
            r3 = p3.add_run(f"Par {vedette['author']}  ·  {vedette.get('date','')}")
            r3.bold = True; r3.font.size = Pt(7)
            r3.font.color.rgb = RGBColor(168,180,207); r3.font.name = "Arial"
        txt = body_text(vedette.get("body",[]))
        if txt:
            p4 = cl.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p4.paragraph_format.line_spacing = Pt(12)
            p4.paragraph_format.space_before = Pt(0); p4.paragraph_format.space_after = Pt(0)
            r4 = p4.add_run(txt[:1200])
            r4.font.size = Pt(9); r4.font.color.rgb = WHITE; r4.font.name = "Times New Roman"
        url = vedette.get("imgUrl","")
        path = BASE / "public" / url.lstrip("/")
        if not path.exists():
            jpg = path.with_suffix(".jpg")
            if jpg.exists(): path = jpg
        cr = tbl.cell(0,1)
        if path.exists() and path.suffix.lower() != ".svg":
            try:
                p = cr.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
                run = p.add_run()
                run.add_picture(str(path), width=Cm(7.5), height=Cm(9.5))
            except Exception:
                set_cell_bg(cr, "0A1E45")
        else:
            set_cell_bg(cr, "0A1E45")
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    if secondaires:
        tbl2 = doc.add_table(rows=1, cols=len(secondaires))
        remove_table_borders(tbl2)
        add_colored_border_top(tbl2, "C8302A", size=10)
        add_insideV_border(tbl2)
        w_each = 18.5 / len(secondaires)
        set_col_widths(tbl2, [w_each] * len(secondaires))
        for i, art in enumerate(secondaires):
            cell = tbl2.cell(0, i)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            add_article(cell, art, title_size=11, show_img=True,
                        img_w_cm=w_each-0.4, img_h_cm=2.8)
    add_footer_line(doc, 1, edition_date)

IMG_W_COL3 = 5.5
IMG_H_COL3 = 3.2
IMG_W_COL2 = 8.5
IMG_H_COL2 = 4.5

def build_standard_page(doc, articles: list, rubrique_label: str,
                         page_num: int, edition_date: datetime, numero: int):
    add_section_break(doc)
    add_inner_header(doc, rubrique_label, edition_date, numero)
    n = len(articles)
    if n == 0:
        doc.add_paragraph(f"Aucun article — {rubrique_label}")
        add_footer_line(doc, page_num, edition_date)
        return
    cols  = 1 if n == 1 else (2 if n == 2 else 3)
    img_w = IMG_W_COL2 if cols == 2 else IMG_W_COL3
    img_h = IMG_H_COL2 if cols == 2 else IMG_H_COL3
    buckets: list = [[] for _ in range(cols)]
    for i, a in enumerate(articles):
        buckets[i % cols].append(a)
    tbl = doc.add_table(rows=1, cols=cols)
    remove_table_borders(tbl)
    add_insideV_border(tbl)
    w_each = 18.5 / cols
    set_col_widths(tbl, [w_each] * cols)
    for i in range(cols):
        cell = tbl.cell(0, i)
        set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
        for j, art in enumerate(buckets[i]):
            add_article(cell, art,
                        title_size=14 if (j==0 and i==0) else (12 if j==0 else 10),
                        show_img=(j == 0),
                        img_w_cm=img_w, img_h_cm=img_h)
            if j < len(buckets[i]) - 1:
                add_rule(cell, "0E2B62", size=6)
    add_footer_line(doc, page_num, edition_date)

def build_page8(doc, articles: list, page_num: int, edition_date: datetime, numero: int):
    add_section_break(doc)
    add_inner_header(doc, "GRANDE INTERVIEW · PORTRAIT · CONTRIBUTION", edition_date, numero)
    if not articles:
        doc.add_paragraph("Espace réservé — Grande interview / Portrait / Contribution")
        add_footer_line(doc, page_num, edition_date)
        return
    art = articles[0]
    tbl = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl)
    add_insideV_border(tbl)
    set_col_widths(tbl, [9.5, 9.0])
    for i in range(2):
        cell = tbl.cell(0, i)
        set_cell_margins(cell, top=60, bottom=60, left=130, right=130)
        if i == 0:
            add_article(cell, art, title_size=16, show_img=False)
        else:
            url = art.get("imgUrl","")
            path = BASE / "public" / url.lstrip("/")
            if not path.exists():
                jpg = path.with_suffix(".jpg")
                if jpg.exists(): path = jpg
            if path.exists() and path.suffix.lower() != ".svg":
                try:
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(6)
                    p.add_run().add_picture(str(path), width=Cm(8.5), height=Cm(6.0))
                except Exception:
                    pass
            txt = body_text(art.get("body",[]))
            words = txt.split()
            half  = len(words) // 2
            suite = " ".join(words[half:])
            if suite:
                p2 = cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p2.paragraph_format.line_spacing = Pt(12)
                r2 = p2.add_run(suite)
                r2.font.size = Pt(9); r2.font.color.rgb = INK; r2.font.name = "Times New Roman"
            if len(articles) > 1:
                add_rule(cell, "0E2B62", 6)
                add_article(cell, articles[1], title_size=11, show_img=False)
    add_footer_line(doc, page_num, edition_date)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--date",   default=datetime.today().strftime("%Y-%m-%d"))
    parser.add_argument("--numero", type=int, default=1)
    parser.add_argument("--out",    default=None)
    args = parser.parse_args()
    edition_date = datetime.strptime(args.date, "%Y-%m-%d")
    numero       = args.numero
    data   = BASE / "src" / "data" / "articles.json"
    outdir = BASE / "public" / "journal"
    outdir.mkdir(exist_ok=True)
    out_path = args.out or str(outdir / f"alheure-{args.date}.docx")
    with open(data, encoding="utf-8") as f:
        all_arts = json.load(f)
    articles = [a for a in all_arts if a.get("status","published") in ("published","",None)]
    def by_rub(*rubriques):
        return [a for a in articles if a.get("rubrique","") in rubriques]
    actualites = by_rub("senegal","politique","economie","monde")
    societe    = by_rub("societe","diaspora","sante","education")
    afrique    = by_rub("afrique")
    sport      = by_rub("sport")
    culture    = by_rub("culture")
    special    = (by_rub("interview","portrait","contribution") or
                  [a for a in articles if a.get("badge") == "longformat"] or
                  articles[-2:])
    doc = Document()
    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(1.0)
        section.bottom_margin = Cm(1.0)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(9)
    mid = max(1, len(actualites)//2)
    build_une(doc, articles, edition_date, numero)
    build_standard_page(doc, actualites[:mid] or articles[:4],  "ACTUALITÉS",   2, edition_date, numero)
    build_standard_page(doc, actualites[mid:] or articles[4:8], "ACTUALITÉS",   3, edition_date, numero)
    build_standard_page(doc, societe  or articles[8:11],        "SOCIÉTÉ",      4, edition_date, numero)
    build_standard_page(doc, afrique  or articles[11:14],       "ACTU AFRIQUE", 5, edition_date, numero)
    build_standard_page(doc, sport    or articles[14:16],       "SPORT",        6, edition_date, numero)
    build_standard_page(doc, culture  or articles[16:19],       "CULTURE",      7, edition_date, numero)
    build_page8(doc, special, 8, edition_date, numero)
    doc.save(out_path)
    print(f"✓ Journal DOCX généré : {out_path}")

if __name__ == "__main__":
    main()
